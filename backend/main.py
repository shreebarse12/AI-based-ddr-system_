"""
DDR Report Generation System — FastAPI Backend
"""

import os
import io
import json
import uuid
import base64
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

from groq import Groq
import pdfplumber
from pypdf import PdfReader
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, Image as RLImage
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# PyMuPDF — optional, graceful fallback if DLL fails on Windows
try:
    import fitz
    FITZ_AVAILABLE = True
except (ImportError, Exception):
    FITZ_AVAILABLE = False

app = FastAPI(
    title="DDR Report Generator API",
    description="AI-powered Detailed Diagnostic Report generation from inspection documents",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve frontend
frontend_path = Path(__file__).parent.parent / "frontend"
if frontend_path.exists():
    app.mount("/static", StaticFiles(directory=str(frontend_path)), name="static")

EXPORT_DIR = Path(__file__).parent.parent / "exports"
EXPORT_DIR.mkdir(exist_ok=True)

# ─── Anthropic client ────────────────────────────────────────────────────────
client = Groq(api_key=os.environ.get("GROQ_API_KEY", ""))

SYSTEM_PROMPT = """You are a senior building diagnostics engineer with 20+ years of experience.
Your task: read raw inspection and thermal report data, then produce a structured DDR (Detailed Diagnostic Report) as JSON.

STRICT RULES:
- NEVER invent facts not in the documents
- If data conflicts between documents → explicitly note the conflict
- If data is missing → use exactly "Not Available"
- Use plain, client-friendly language (no excessive jargon)
- Be specific: mention exact locations, temperatures, measurements when available

Return ONLY valid JSON with this exact structure:
{
  "property_summary": "2-4 sentence executive overview of all key issues",
  "areas": [
    {
      "name": "Exact area name (e.g. Roof Section A, East Basement Wall)",
      "severity": "High | Medium | Low",
      "observations": "Detailed observation combining inspection + thermal findings",
      "thermal_finding": "Specific temperature data or thermal anomaly, or 'Not Available'",
      "image_label": "Short label for image placement, or null if no image expected"
    }
  ],
  "root_causes": [
    {
      "issue": "Issue title",
      "cause": "Probable root cause with reasoning"
    }
  ],
  "severity_assessment": [
    {
      "area": "Area name",
      "severity": "High | Medium | Low",
      "reasoning": "Why this severity level was assigned"
    }
  ],
  "recommended_actions": [
    "Specific, actionable recommendation"
  ],
  "additional_notes": "Important context, warnings, or disclaimers. 'Not Available' if none.",
  "missing_info": ["List of specific missing or unclear data points"],
  "conflicts": ["List of any conflicting data between the two documents, or empty array"]
}"""


# ─── Text extraction ─────────────────────────────────────────────────────────

def extract_text_from_pdf(data: bytes) -> str:
    """Extract text from PDF using pdfplumber for best layout preservation."""
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                # Also try extracting tables
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            if row:
                                text += "\n" + " | ".join(str(c) for c in row if c)
                pages.append(f"[Page {i+1}]\n{text}")
            return "\n\n".join(pages)
    except Exception as e:
        # Fallback to pypdf
        try:
            reader = PdfReader(io.BytesIO(data))
            return "\n\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            return f"[Could not extract text: {e}]"


def extract_text_from_file(data: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(data)
    elif ext in ("txt", "md"):
        return data.decode("utf-8", errors="replace")
    elif ext in ("png", "jpg", "jpeg"):
        return "[IMAGE FILE — visual content only, no extractable text]"
    else:
        return data.decode("utf-8", errors="replace")


def extract_images_from_pdf(data: bytes, label: str, max_pages: int = 6) -> list[dict]:
    """Render PDF pages as base64 images. Uses PyMuPDF if available, else pdf2image, else skips."""
    images = []

    # Method 1: PyMuPDF (best, but has Windows DLL issues)
    if FITZ_AVAILABLE:
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            for i in range(min(len(doc), max_pages)):
                page = doc[i]
                mat = fitz.Matrix(1.5, 1.5)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("jpeg")
                b64 = base64.b64encode(img_data).decode()
                images.append({
                    "src": f"data:image/jpeg;base64,{b64}",
                    "label": f"{label} — Page {i+1}",
                    "page": i + 1
                })
            doc.close()
            return images
        except Exception as e:
            print(f"PyMuPDF error: {e}, trying fallback...")

    # Method 2: pdf2image (requires poppler on PATH)
    try:
        from pdf2image import convert_from_bytes
        pages = convert_from_bytes(data, dpi=120, first_page=1, last_page=max_pages)
        for i, page_img in enumerate(pages):
            buf = io.BytesIO()
            page_img.save(buf, format="JPEG", quality=75)
            b64 = base64.b64encode(buf.getvalue()).decode()
            images.append({
                "src": f"data:image/jpeg;base64,{b64}",
                "label": f"{label} — Page {i+1}",
                "page": i + 1
            })
        return images
    except Exception as e:
        print(f"pdf2image fallback error: {e}")

    # Method 3: No image extraction available — report shows placeholder
    print(f"Warning: Image extraction unavailable for {label}")
    return images


# ─── AI reasoning ────────────────────────────────────────────────────────────

def call_claude(inspection_text: str, thermal_text: str) -> dict:
    """Call Groq API and parse structured DDR JSON."""
    user_message = f"""INSPECTION REPORT:
{inspection_text[:6000]}

THERMAL REPORT:
{thermal_text[:4000]}

Generate the DDR JSON now. Return ONLY valid JSON, no preamble, no markdown fences."""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        max_tokens=2000,
        temperature=0.1,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_message}
        ]
    )

    raw = response.choices[0].message.content or ""
    # Strip markdown fences if present
    clean = raw.strip()
    if clean.startswith("```"):
        clean = clean.split("```")[1]
        if clean.startswith("json"):
            clean = clean[4:]
    clean = clean.strip().rstrip("```").strip()

    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        import re
        match = re.search(r'\{[\s\S]*\}', clean)
        if match:
            return json.loads(match.group())
        raise ValueError("Could not parse AI response as JSON. Raw: " + raw[:200])


# ─── PDF export ───────────────────────────────────────────────────────────────

def build_pdf_report(report: dict, images: list[dict], out_path: str):
    """Generate a professional PDF report using ReportLab."""
    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2*cm
    )

    W = A4[0] - 4*cm  # usable width

    styles = getSampleStyleSheet()
    DARK = colors.HexColor("#1a1f2e")
    ACCENT = colors.HexColor("#e8a020")
    MUTED = colors.HexColor("#6b7280")
    HIGH = colors.HexColor("#dc2626")
    MED = colors.HexColor("#d97706")
    LOW = colors.HexColor("#16a34a")

    title_style = ParagraphStyle("DDRTitle", parent=styles["Title"],
        fontSize=22, textColor=DARK, spaceAfter=4, fontName="Helvetica-Bold")
    sub_style = ParagraphStyle("DDRSub", parent=styles["Normal"],
        fontSize=9, textColor=MUTED, spaceAfter=20)
    heading_style = ParagraphStyle("DDRHeading", parent=styles["Heading1"],
        fontSize=13, textColor=DARK, fontName="Helvetica-Bold",
        spaceBefore=18, spaceAfter=8,
        borderPad=6, leftIndent=0)
    body_style = ParagraphStyle("DDRBody", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#374151"),
        leading=15, spaceAfter=6)
    label_style = ParagraphStyle("DDRLabel", parent=styles["Normal"],
        fontSize=8, textColor=MUTED, fontName="Helvetica-Oblique")

    def sev_color(s):
        s = (s or "").lower()
        return HIGH if s == "high" else MED if s == "medium" else LOW

    def section_title(num, title):
        return [
            HRFlowable(width=W, thickness=2, color=ACCENT, spaceAfter=6),
            Paragraph(f"<font color='#e8a020'>{num}</font>  {title}", heading_style),
        ]

    story = []

    # ── Cover ──
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("DETAILED DIAGNOSTIC REPORT", title_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}  ·  AI-Synthesised", sub_style))
    story.append(HRFlowable(width=W, thickness=3, color=ACCENT, spaceAfter=20))
    story.append(Spacer(1, 0.3*cm))

    # ── 1. Property Summary ──
    story += section_title("01", "Property Issue Summary")
    story.append(Paragraph(report.get("property_summary", "Not Available"), body_style))
    story.append(Spacer(1, 0.3*cm))

    # ── 2. Area-wise Observations ──
    story += section_title("02", "Area-wise Observations")
    img_map = {img["label"]: img["src"] for img in images}

    for area in report.get("areas", []):
        sev = area.get("severity", "Medium")
        sc = sev_color(sev)
        name_para = Paragraph(
            f'<b>{area.get("name","")}</b>  <font color="{sc.hexval()}" size="8">[{sev.upper()}]</font>',
            body_style
        )
        obs_para = Paragraph(area.get("observations", ""), body_style)
        thermal = area.get("thermal_finding", "Not Available")
        thermal_para = Paragraph(
            f'<font color="#2563eb"><b>Thermal:</b> {thermal}</font>', body_style
        )
        tdata = [[name_para], [obs_para], [thermal_para]]
        tbl = Table(tdata, colWidths=[W])
        tbl.setStyle(TableStyle([
            ("BOX", (0,0), (-1,-1), 0.5, colors.HexColor("#d1d5db")),
            ("BACKGROUND", (0,0), (0,0), colors.HexColor("#f9fafb")),
            ("LEFTPADDING", (0,0), (-1,-1), 10),
            ("RIGHTPADDING", (0,0), (-1,-1), 10),
            ("TOPPADDING", (0,0), (-1,-1), 8),
            ("BOTTOMPADDING", (0,0), (-1,-1), 8),
        ]))
        story.append(tbl)

        # Attach image if available
        lbl = area.get("image_label")
        matched_img = None
        if lbl:
            for img in images:
                if lbl.lower() in img["label"].lower() or img["label"].lower() in lbl.lower():
                    matched_img = img
                    break
        if not matched_img and images:
            # Attach first unused image heuristically
            matched_img = images[0] if images else None

        if matched_img:
            try:
                b64data = matched_img["src"].split(",", 1)[1]
                img_bytes = base64.b64decode(b64data)
                img_buf = io.BytesIO(img_bytes)
                rl_img = RLImage(img_buf, width=W*0.6, height=W*0.35)
                story.append(rl_img)
                story.append(Paragraph(matched_img["label"], label_style))
            except Exception:
                story.append(Paragraph("[ Image could not be rendered ]", label_style))
        else:
            story.append(Paragraph("[ Image Not Available ]", label_style))

        story.append(Spacer(1, 0.25*cm))

    # ── 3. Root Cause ──
    story += section_title("03", "Probable Root Cause")
    for i, rc in enumerate(report.get("root_causes", []), 1):
        story.append(Paragraph(f"<b>{i}. {rc.get('issue','')}</b>", body_style))
        story.append(Paragraph(rc.get("cause",""), body_style))
    story.append(Spacer(1, 0.2*cm))

    # ── 4. Severity Assessment ──
    story += section_title("04", "Severity Assessment")
    sev_rows = [["Area", "Severity", "Reasoning"]]
    for s in report.get("severity_assessment", []):
        sev_rows.append([s.get("area",""), s.get("severity",""), s.get("reasoning","")])
    if len(sev_rows) > 1:
        sev_tbl = Table(sev_rows, colWidths=[W*0.25, W*0.12, W*0.63])
        sev_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), DARK),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f9fafb")]),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#e5e7eb")),
            ("LEFTPADDING", (0,0), (-1,-1), 8),
            ("RIGHTPADDING", (0,0), (-1,-1), 8),
            ("TOPPADDING", (0,0), (-1,-1), 6),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
        ]))
        story.append(sev_tbl)
    story.append(Spacer(1, 0.3*cm))

    # ── 5. Recommended Actions ──
    story += section_title("05", "Recommended Actions")
    for i, action in enumerate(report.get("recommended_actions", []), 1):
        story.append(Paragraph(f"{i}.  {action}", body_style))
    story.append(Spacer(1, 0.2*cm))

    # ── 6. Additional Notes ──
    story += section_title("06", "Additional Notes")
    story.append(Paragraph(report.get("additional_notes", "Not Available"), body_style))
    story.append(Spacer(1, 0.2*cm))

    # ── 7. Missing / Conflicts ──
    story += section_title("07", "Missing or Unclear Information")
    missing = report.get("missing_info", [])
    conflicts = report.get("conflicts", [])
    if not missing and not conflicts:
        story.append(Paragraph("No missing information identified.", body_style))
    for m in missing:
        story.append(Paragraph(f"⚠  {m}", body_style))
    for c in conflicts:
        story.append(Paragraph(f"⚡  CONFLICT: {c}", body_style))

    doc.build(story)


# ─── DOCX export ──────────────────────────────────────────────────────────────

def build_docx_report(report: dict, images: list[dict], out_path: str):
    """Generate a professional Word document."""
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("Detailed Diagnostic Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1f, 0x2e)
    title.runs[0].font.size = Pt(24)

    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}  ·  AI-Synthesised")
    meta.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    meta.runs[0].font.size = Pt(9)

    def add_hr(doc):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'e8a020')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section(num, title_text):
        add_hr(doc)
        h = doc.add_heading(f"{num}  {title_text}", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1f, 0x2e)
        h.runs[0].font.size = Pt(13)

    def sev_color_rgb(s):
        s = (s or "").lower()
        if s == "high": return RGBColor(0xdc, 0x26, 0x26)
        if s == "medium": return RGBColor(0xd9, 0x77, 0x06)
        return RGBColor(0x16, 0xa3, 0x4a)

    # 1. Property Summary
    add_section("01", "Property Issue Summary")
    doc.add_paragraph(report.get("property_summary", "Not Available"))

    # 2. Area-wise Observations
    add_section("02", "Area-wise Observations")
    for area in report.get("areas", []):
        sev = area.get("severity", "Medium")
        p = doc.add_paragraph()
        run = p.add_run(area.get("name", ""))
        run.bold = True
        run.font.size = Pt(11)
        sev_run = p.add_run(f"  [{sev.upper()}]")
        sev_run.font.color.rgb = sev_color_rgb(sev)
        sev_run.font.size = Pt(9)

        doc.add_paragraph(area.get("observations", ""))

        thermal = area.get("thermal_finding", "Not Available")
        tp = doc.add_paragraph()
        tr = tp.add_run(f"Thermal Finding: {thermal}")
        tr.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
        tr.font.size = Pt(9)

        # Add first available image per area
        if images:
            try:
                img = images[0]
                b64data = img["src"].split(",", 1)[1]
                img_bytes = base64.b64decode(b64data)
                buf = io.BytesIO(img_bytes)
                doc.add_picture(buf, width=Inches(4))
                cap = doc.add_paragraph(img["label"])
                cap.runs[0].font.size = Pt(8)
                cap.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
                cap.runs[0].font.italic = True
            except Exception:
                doc.add_paragraph("[Image Not Available]")
        else:
            doc.add_paragraph("[Image Not Available]")

    # 3. Root Cause
    add_section("03", "Probable Root Cause")
    for i, rc in enumerate(report.get("root_causes", []), 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{rc.get('issue','')}: ")
        run.bold = True
        p.add_run(rc.get("cause", ""))

    # 4. Severity Assessment
    add_section("04", "Severity Assessment")
    sev_data = report.get("severity_assessment", [])
    if sev_data:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Area"
        hdr[1].text = "Severity"
        hdr[2].text = "Reasoning"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), '1a1f2e')
            cell._tc.get_or_add_tcPr().append(shading)
        for s in sev_data:
            row = tbl.add_row().cells
            row[0].text = s.get("area", "")
            row[1].text = s.get("severity", "")
            row[1].paragraphs[0].runs[0].font.color.rgb = sev_color_rgb(s.get("severity",""))
            row[2].text = s.get("reasoning", "")

    # 5. Recommended Actions
    add_section("05", "Recommended Actions")
    for action in report.get("recommended_actions", []):
        doc.add_paragraph(action, style="List Number")

    # 6. Additional Notes
    add_section("06", "Additional Notes")
    doc.add_paragraph(report.get("additional_notes", "Not Available"))

    # 7. Missing Info
    add_section("07", "Missing or Unclear Information")
    missing = report.get("missing_info", [])
    conflicts = report.get("conflicts", [])
    if not missing and not conflicts:
        doc.add_paragraph("No missing information identified.")
    for m in missing:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(m)
    for c in conflicts:
        p = doc.add_paragraph()
        r = p.add_run(f"CONFLICT: {c}")
        r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

    doc.save(out_path)


# ─── API Endpoints ────────────────────────────────────────────────────────────

@app.get("/")
async def root():
    index = frontend_path / "index.html"
    if index.exists():
        return FileResponse(str(index))
    return {"message": "DDR Report Generator API", "docs": "/docs"}


@app.post("/api/generate")
async def generate_report(
    inspection: UploadFile = File(...),
    thermal: UploadFile = File(...)
):
    """
    Accept inspection + thermal documents, run AI analysis, return structured DDR JSON + images.
    """
    insp_data = await inspection.read()
    therm_data = await thermal.read()

    insp_text = extract_text_from_file(insp_data, inspection.filename)
    therm_text = extract_text_from_file(therm_data, thermal.filename)

    if not insp_text.strip() and not therm_text.strip():
        raise HTTPException(400, "Could not extract any text from the uploaded files.")

    # Extract images
    images = []
    if inspection.filename.lower().endswith(".pdf"):
        images += extract_images_from_pdf(insp_data, "Inspection")
    elif inspection.filename.lower().endswith((".png", ".jpg", ".jpeg")):
        b64 = base64.b64encode(insp_data).decode()
        ext = inspection.filename.rsplit(".", 1)[-1].lower()
        images.append({"src": f"data:image/{ext};base64,{b64}", "label": "Inspection Image", "page": 1})

    if thermal.filename.lower().endswith(".pdf"):
        images += extract_images_from_pdf(therm_data, "Thermal")
    elif thermal.filename.lower().endswith((".png", ".jpg", ".jpeg")):
        b64 = base64.b64encode(therm_data).decode()
        ext = thermal.filename.rsplit(".", 1)[-1].lower()
        images.append({"src": f"data:image/{ext};base64,{b64}", "label": "Thermal Image", "page": 1})

    try:
        report_data = call_claude(insp_text, therm_text)
    except Exception as e:
        raise HTTPException(500, f"AI analysis failed: {str(e)}")

    report_id = str(uuid.uuid4())[:8]
    return JSONResponse({
        "report_id": report_id,
        "generated_at": datetime.now().isoformat(),
        "report": report_data,
        "images": images
    })


@app.post("/api/export/pdf")
async def export_pdf(payload: dict):
    """Generate and return a PDF from report data."""
    report = payload.get("report", {})
    images = payload.get("images", [])
    report_id = payload.get("report_id", "ddr")

    out_path = str(EXPORT_DIR / f"DDR_{report_id}.pdf")
    try:
        build_pdf_report(report, images, out_path)
    except Exception as e:
        raise HTTPException(500, f"PDF generation failed: {str(e)}")

    return FileResponse(
        out_path,
        media_type="application/pdf",
        filename=f"DDR_Report_{report_id}.pdf"
    )


@app.post("/api/export/docx")
async def export_docx(payload: dict):
    """Generate and return a DOCX from report data."""
    report = payload.get("report", {})
    images = payload.get("images", [])
    report_id = payload.get("report_id", "ddr")

    out_path = str(EXPORT_DIR / f"DDR_{report_id}.docx")
    try:
        build_docx_report(report, images, out_path)
    except Exception as e:
        raise HTTPException(500, f"DOCX generation failed: {str(e)}")

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"DDR_Report_{report_id}.docx"
    )


@app.get("/health")
async def health():
    return {"status": "ok", "version": "1.0.0"}